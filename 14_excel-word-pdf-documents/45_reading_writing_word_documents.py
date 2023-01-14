import docx

# Storing document as a variable: -
d = docx.Document("demo.docx")

# Printing a paragraph: -
print(d.paragraphs[0].text)
print(d.paragraphs[1].text)

p = d.paragraphs[1]

# RUN objects are whenever a change in style occurs in text: -
print(p.runs)
print(p.runs[0].text) 
print(p.runs[1].text) 
print(p.runs[2].text) 
print(p.runs[3].text) 

# You can check if a run object has a member variable set to: -
print(p.runs[1].bold) 
print(p.runs[3].italic) 

# You can then assign member varaibles to a run objects and change content: -
p.runs[3].underline = True
p.runs[3].text = "italic and underlined."

# Saving a new document from loaded memory: -
d.save("demo2.docx")

# Changing style: -
p.style = "Title"
d.save("demo3.docx")

# Creating a brand new word document: -
d = docx.Document()

# Add a new paragraph: -
d.add_paragraph("Hello, this is a paragraph.")
d.add_paragraph("This is another paragraph.")
d.save("demo4.docx")

# Add a new run and make it bold: -
p = d.paragraphs[0]
p.add_run("This is a new run.")
p.runs[1].bold = True
d.save("demo5.docx")

# You can only add paragraphs and runs to the END of a document.
# So to edit a word document and add some data inbetween you'd have to make a new document,
# use a for loop to import up to where you want you edit the document,
# add in your new content, then continue adding in the rest of the original test.

# A simple function to copy all the data from one document and print it in the terminal: -
# (Which of course strips the formatting so a new line has been added inbetween each string)
def get_text(filename):
    doc = docx.Document(filename)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return "\n".join(full_text)
print(get_text("demo.docx"))
